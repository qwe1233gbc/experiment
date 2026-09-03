#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Pilot17 修复 - 阶段2：17题Word→JSON→上下文逐证据槽审计
直接读取原始Word，建立逐证据槽追踪表
"""
import re
import json
import csv
import hashlib
from pathlib import Path
from collections import defaultdict

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️  python-docx 未安装，尝试安装...")
    import subprocess
    subprocess.run(["pip", "install", "python-docx"], capture_output=True)
    from docx import Document
    HAS_DOCX = True

EXPERIMENT_ROOT = Path(r"E:\实验文件整理_按论文逻辑\实验")
REPAIR_DIR = EXPERIMENT_ROOT / "07_results_v2" / "pilot17_repair_20260903"

# 引入规范化搜索
import sys
sys.path.insert(0, str(REPAIR_DIR))
from evidence_search_utils import normalize_for_evidence_search, contains_standard, find_standard_occurrences

# ============ Word读取 ============

def extract_word_text(docx_path):
    """从Word文档提取所有文本（段落+表格），带位置标记"""
    doc = Document(docx_path)
    
    all_text_parts = []
    table_count = 0
    para_count = 0
    
    # 段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text_parts.append({
                "type": "paragraph",
                "index": i,
                "style": para.style.name if para.style else "",
                "text": text,
            })
            para_count += 1
    
    # 表格
    for t_idx, table in enumerate(doc.tables):
        table_text_parts = []
        for r_idx, row in enumerate(table.rows):
            row_cells = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
            table_text_parts.append(row_cells)
        
        # 表格转为文本
        flat_text = "\n".join([" | ".join(row) for row in table_text_parts])
        
        all_text_parts.append({
            "type": "table",
            "index": t_idx,
            "text": flat_text,
            "row_count": len(table_text_parts),
            "col_count": len(table_text_parts[0]) if table_text_parts else 0,
        })
        table_count += 1
    
    # 合并全文
    full_text = "\n".join([p["text"] for p in all_text_parts])
    
    return {
        "paragraph_count": para_count,
        "table_count": table_count,
        "full_text": full_text,
        "parts": all_text_parts,
        "total_chars": len(full_text),
    }

# ============ 17题证据槽定义 ============

def get_question_evidence_slots():
    """定义17题的必需证据槽"""
    slots = {
        # 环保投资比例核算（4题，模板相同）
        "NEW_PL001_invest_ratio": {
            "project": "PL001",
            "task_type": "环保投资比例核算",
            "slots": [
                {"id": "E1", "fact": "总投资额", "keywords": ["总投资", "总投资额"], "internal": True},
                {"id": "E2", "fact": "环保投资额", "keywords": ["环保投资", "环保投资额"], "internal": True},
                {"id": "E3", "fact": "环保投资比例填报值", "keywords": ["环保投资比例", "占比", "比例"], "internal": True},
            ],
            "needs_web": False,
            "needs_rag": False,
        },
        "NEW_PL006_invest_ratio": {
            "project": "PL006",
            "task_type": "环保投资比例核算",
            "slots": [
                {"id": "E1", "fact": "总投资额", "keywords": ["总投资", "总投资额"], "internal": True},
                {"id": "E2", "fact": "环保投资额", "keywords": ["环保投资", "环保投资额"], "internal": True},
                {"id": "E3", "fact": "环保投资比例填报值", "keywords": ["环保投资比例", "占比", "比例"], "internal": True},
            ],
            "needs_web": False,
            "needs_rag": False,
        },
        "NEW_PL010_invest_ratio": {
            "project": "PL010",
            "task_type": "环保投资比例核算",
            "slots": [
                {"id": "E1", "fact": "总投资额(15000万元)", "keywords": ["总投资", "15000", "总投资额"], "internal": True},
                {"id": "E2", "fact": "环保投资额(174.5万元)", "keywords": ["环保投资", "174.5"], "internal": True},
                {"id": "E3", "fact": "环保投资比例填报值(1.1%)", "keywords": ["1.1%", "环保投资比例", "占比"], "internal": True},
            ],
            "needs_web": False,
            "needs_rag": False,
        },
        "NEW_PL015_invest_ratio": {
            "project": "PL015",
            "task_type": "环保投资比例核算",
            "slots": [
                {"id": "E1", "fact": "总投资额", "keywords": ["总投资", "总投资额"], "internal": True},
                {"id": "E2", "fact": "环保投资额", "keywords": ["环保投资", "环保投资额"], "internal": True},
                {"id": "E3", "fact": "环保投资比例填报值", "keywords": ["环保投资比例", "占比", "比例"], "internal": True},
            ],
            "needs_web": False,
            "needs_rag": False,
        },
        # 国民经济行业分类
        "PL002_V01_Q01": {
            "project": "PL002",
            "task_type": "国民经济行业分类",
            "slots": [
                {"id": "E1", "fact": "产品描述", "keywords": ["产品", "改性塑料", "塑料粒"], "internal": True},
                {"id": "E2", "fact": "原辅料描述", "keywords": ["原料", "玻璃纤维", "色母"], "internal": True},
                {"id": "E3", "fact": "生产工艺", "keywords": ["双螺杆", "挤出", "切粒", "混料"], "internal": True},
                {"id": "E4", "fact": "行业类别填报", "keywords": ["C2929", "行业类别", "国民经济行业"], "internal": True},
                {"id": "E5", "fact": "行业分类标准(GB/T 4754)", "keywords": ["GB/T 4754", "行业分类"], "internal": False, "web_only": True},
            ],
            "needs_web": True,
            "needs_rag": False,
        },
        "PL004_V01_Q01": {
            "project": "PL004",
            "task_type": "国民经济行业分类",
            "slots": [
                {"id": "E1", "fact": "产品描述", "keywords": ["产品", "SEBS", "塑料粒"], "internal": True},
                {"id": "E2", "fact": "原辅料描述", "keywords": ["原料", "投料"], "internal": True},
                {"id": "E3", "fact": "生产工艺", "keywords": ["双螺杆", "挤出", "切粒"], "internal": True},
                {"id": "E4", "fact": "行业类别填报", "keywords": ["C2929", "行业类别", "国民经济行业"], "internal": True},
                {"id": "E5", "fact": "行业分类标准", "keywords": ["GB/T 4754", "行业分类"], "internal": False, "web_only": True},
            ],
            "needs_web": True,
            "needs_rag": False,
        },
        # 固体废物控制标准
        "PL001_Emission_固体": {
            "project": "PL001",
            "task_type": "固体废物控制标准",
            "slots": [
                {"id": "E1", "fact": "一般固废暂存方式", "keywords": ["一般工业固体废物", "库房暂存", "暂存"], "internal": True},
                {"id": "E2", "fact": "一般固废标准GB18599", "keywords": ["GB18599-2020", "GB 18599", "一般工业固体废物贮存"], "internal": True},
                {"id": "E3", "fact": "危废标准GB18597", "keywords": ["GB18597-2023", "GB 18597", "危险废物贮存"], "internal": True},
                {"id": "E4", "fact": "固废鉴别GB34330", "keywords": ["GB34330-2017", "GB 34330", "固体废物鉴别"], "internal": True},
                {"id": "E5", "fact": "国家危险废物名录", "keywords": ["国家危险废物名录", "危废名录"], "internal": True},
                {"id": "E6", "fact": "危废识别与类别代码", "keywords": ["HW08", "HW49", "危险废物", "废物代码"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
        "PL005_Emission_固体": {
            "project": "PL005",
            "task_type": "固体废物控制标准",
            "slots": [
                {"id": "E1", "fact": "一般固废暂存方式", "keywords": ["一般工业固体废物", "仓库", "暂存"], "internal": True},
                {"id": "E2", "fact": "一般固废标准GB18599", "keywords": ["GB18599-2020", "GB 18599", "一般工业固体废物贮存"], "internal": True},
                {"id": "E3", "fact": "危废标准GB18597", "keywords": ["GB18597-2023", "GB 18597", "危险废物贮存"], "internal": True},
                {"id": "E4", "fact": "固废分类GB/T39198", "keywords": ["GB/T39198-2020", "GB/T 39198"], "internal": True},
                {"id": "E5", "fact": "国家危险废物名录", "keywords": ["国家危险废物名录", "危废名录"], "internal": True},
                {"id": "E6", "fact": "三防要求(防渗漏防雨淋防扬尘)", "keywords": ["防渗漏", "防雨淋", "防扬尘", "三防"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
        # VOCs源强与总量一致性
        "PL008_VOCSTotal_Q01": {
            "project": "PL008",
            "task_type": "VOCs源强与总量数值一致性",
            "slots": [
                {"id": "E1", "fact": "源强表有组织排放量", "keywords": ["有组织", "VOCs", "排放量"], "internal": True},
                {"id": "E2", "fact": "源强表无组织排放量", "keywords": ["无组织", "VOCs", "排放量"], "internal": True},
                {"id": "E3", "fact": "源强表明细合计", "keywords": ["合计", "总计", "源强"], "internal": True},
                {"id": "E4", "fact": "总量控制章节数据", "keywords": ["总量控制", "VOCs", "排放量"], "internal": True},
                {"id": "E5", "fact": "关键数值0.9089/0.90867", "keywords": ["0.9089", "0.90867", "3.1802"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
        # VOCs治理措施一致性
        "PL008_VOCSMeasure_Q01": {
            "project": "PL008",
            "task_type": "VOCs治理措施一致性",
            "slots": [
                {"id": "E1", "fact": "工程分析章节VOCs治理设施", "keywords": ["工程分析", "VOCs", "活性炭", "吸附"], "internal": True},
                {"id": "E2", "fact": "运营期措施章节VOCs治理", "keywords": ["运营期", "环境保护措施", "VOCs", "活性炭"], "internal": True},
                {"id": "E3", "fact": "监督检查清单VOCs设施", "keywords": ["监督检查", "VOCs", "活性炭", "排气筒"], "internal": True},
                {"id": "E4", "fact": "排气筒编号FQ-03672", "keywords": ["FQ-03672", "排气筒"], "internal": True},
                {"id": "E5", "fact": "水帘柜+干式过滤(颗粒物)", "keywords": ["水帘柜", "干式过滤", "粉尘", "颗粒物"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": False,
        },
        # VOCs总量控制与一致性
        "PL010_VOCSTotal_Q01": {
            "project": "PL010",
            "task_type": "VOCs总量控制与一致性",
            "slots": [
                {"id": "E1", "fact": "VOCs有组织排放量(0.194)", "keywords": ["有组织", "0.194", "VOCs"], "internal": True},
                {"id": "E2", "fact": "VOCs无组织排放量(0.047)", "keywords": ["无组织", "0.047", "VOCs"], "internal": True},
                {"id": "E3", "fact": "VOCs总排放量(0.211)", "keywords": ["0.211", "总排放量", "总量控制"], "internal": True},
                {"id": "E4", "fact": "总量控制章节", "keywords": ["总量控制", "VOCs"], "internal": True},
                {"id": "E5", "fact": "源强核算表数据", "keywords": ["源强", "核算", "VOCs"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
        # 生活污水量核算
        "NEW_PL006_living_wastewater": {
            "project": "PL006",
            "task_type": "生活污水量核算",
            "slots": [
                {"id": "E1", "fact": "劳动定员", "keywords": ["劳动定员", "员工人数", "职工"], "internal": True},
                {"id": "E2", "fact": "年工作天数", "keywords": ["工作天数", "年工作日", "300天"], "internal": True},
                {"id": "E3", "fact": "食宿情况(食堂宿舍)", "keywords": ["食堂", "宿舍", "食宿"], "internal": True},
                {"id": "E4", "fact": "生活用水量/定额", "keywords": ["生活用水量", "用水定额", "用水量"], "internal": True},
                {"id": "E5", "fact": "生活污水排放量", "keywords": ["生活污水", "排放量", "废水"], "internal": True},
                {"id": "E6", "fact": "排放系数0.9", "keywords": ["排放系数", "0.9"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
        # 纯水制备水量核算
        "NEW_PL007_ro_water": {
            "project": "PL007",
            "task_type": "纯水制备水量核算",
            "slots": [
                {"id": "E1", "fact": "纯水用量(700t/a)", "keywords": ["纯水", "700", "用水量"], "internal": True},
                {"id": "E2", "fact": "纯水制备率(80%)", "keywords": ["制备率", "80%", "回收率"], "internal": True},
                {"id": "E3", "fact": "浓水排放量", "keywords": ["浓水", "排放量", "RO", "反渗透"], "internal": True},
            ],
            "needs_web": False,
            "needs_rag": False,
        },
        # 废气收集效率
        "PL007_CaptureEfficiency_Q01": {
            "project": "PL007",
            "task_type": "废气收集效率",
            "slots": [
                {"id": "E1", "fact": "收集形式(设备直连/集气罩)", "keywords": ["收集", "集气", "分散缸", "直连"], "internal": True},
                {"id": "E2", "fact": "收集效率取值(90%)", "keywords": ["收集效率", "90%"], "internal": True},
                {"id": "E3", "fact": "广东省VOCs减排核算方法", "keywords": ["减排量核算", "收集效率参考", "广东省"], "internal": False, "rag_only": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
        # 废气设计风量
        "PL008_DesignAirflow_Q01": {
            "project": "PL008",
            "task_type": "废气设计风量",
            "slots": [
                {"id": "E1", "fact": "理论排气量", "keywords": ["理论排气量", "排风量"], "internal": True},
                {"id": "E2", "fact": "设计风量(20000m³/h)", "keywords": ["设计风量", "20000", "m³/h"], "internal": True},
                {"id": "E3", "fact": "收集点数量", "keywords": ["收集点", "排气罩", "集气罩"], "internal": True},
                {"id": "E4", "fact": "治理设施处理能力", "keywords": ["处理能力", "风量", "活性炭"], "internal": True},
            ],
            "needs_web": False,
            "needs_rag": True,
        },
        # 危险废物识别
        "PL013_HazardousWaste_Q01": {
            "project": "PL013",
            "task_type": "危险废物识别",
            "slots": [
                {"id": "E1", "fact": "废机油/HW08/900-249-08", "keywords": ["废机油", "HW08", "900-249-08"], "internal": True},
                {"id": "E2", "fact": "废油桶/HW49", "keywords": ["废油桶", "HW49"], "internal": True},
                {"id": "E3", "fact": "其他危废种类", "keywords": ["危险废物", "废液压油", "废线路板"], "internal": True},
                {"id": "E4", "fact": "危废产生环节匹配", "keywords": ["产生", "设备维护", "废气治理"], "internal": True},
                {"id": "E5", "fact": "国家危险废物名录", "keywords": ["国家危险废物名录", "危废名录"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
        # 废气收集形式与理论排气量
        "PL014_CaptureAirflow_Q01": {
            "project": "PL014",
            "task_type": "废气收集形式与理论排气量",
            "slots": [
                {"id": "E1", "fact": "收集形式(密闭/集气罩)", "keywords": ["集气罩", "密闭", "收集形式"], "internal": True},
                {"id": "E2", "fact": "控制风速法参数", "keywords": ["控制风速", "风速", "m/s"], "internal": True},
                {"id": "E3", "fact": "全面通风量计算(12次/h)", "keywords": ["全面通风", "12次/h", "换气次数"], "internal": True},
                {"id": "E4", "fact": "空间体积数据", "keywords": ["体积", "面积", "高度"], "internal": True},
                {"id": "E5", "fact": "理论排气量计算结果", "keywords": ["理论排气量", "风量", "m³/h"], "internal": True},
            ],
            "needs_web": True,
            "needs_rag": True,
        },
    }
    return slots

# ============ 项目文件映射 ============

PROJECT_FILE_MAP = {
    "PL001": "PL001_佛山市亮正新材料有限公司新建项目",
    "PL002": "PL002_佛山市顺德区启卓工程塑料实业有限公司搬迁项目",
    "PL004": "PL004_恒励新材料科技(佛山)有限公司新建项目",
    "PL005": "PL005_佛山市润特龙清洁用品有限公司新建项目",
    "PL006": "PL006_佛山市顺德区乐鲸科技有限公司年产五金拉手300万个新建项目",
    "PL007": "PL007_007_广东云晟新材料有限公司年产水性涂料2400吨新建项目",
    "PL008": "PL008_佛山市顺德区亚马逊涂料有限公司迁扩建项目",
    "PL010": "PL010_悍高集团股份有限公司功能拉篮车间搬迁扩建项目",
    "PL013": "PL013_佛山市顺德区百洛电器有限公司迁扩建项目",
    "PL014": "PL014_佛山市顺德区蓝顿涂料有限公司迁建项目",
    "PL015": "PL015_佛山市镪优金属制品有限公司迁扩建项目",
}

def search_in_text(text, keywords, context_chars=80):
    """在文本中搜索关键词，返回命中详情"""
    hits = []
    for kw in keywords:
        if not kw:
            continue
        # 先尝试精确匹配
        pos = 0
        kw_lower = kw.lower()
        text_lower = text.lower()
        while True:
            pos = text_lower.find(kw_lower, pos)
            if pos == -1:
                break
            start = max(0, pos - context_chars)
            end = min(len(text), pos + len(kw) + context_chars)
            snippet = text[start:end].replace('\n', ' ')
            snippet = re.sub(r'\s+', ' ', snippet).strip()
            hits.append({
                "keyword": kw,
                "char_pos": pos,
                "snippet": snippet[:200],
                "match_type": "exact",
            })
            pos += len(kw)
    return hits

def search_standard_in_text(text, standard_id):
    """搜索标准编号（使用规范化方法）"""
    if contains_standard(text, standard_id):
        occurrences = find_standard_occurrences(text, standard_id, context_chars=60)
        return occurrences
    return []

def main():
    print("=" * 70)
    print("Pilot17 修复 - 阶段2：17题Word→JSON逐证据槽审计")
    print("=" * 70)
    
    word_dir = EXPERIMENT_ROOT / "09_input_reports" / "原始Word版"
    json_dir = EXPERIMENT_ROOT / "09_input_reports"
    
    slots_def = get_question_evidence_slots()
    
    # 先检查Word文件
    print("\n[1] 读取Word原件...")
    word_data = {}
    for proj_id, base_name in PROJECT_FILE_MAP.items():
        word_path = word_dir / f"{base_name}_01_核查材料.docx"
        # 可能文件名略有不同，尝试匹配
        if not word_path.exists():
            # 查找匹配的文件
            matches = list(word_dir.glob(f"{proj_id}*.docx"))
            if matches:
                word_path = matches[0]
            else:
                print(f"  ⚠️  {proj_id}: 找不到Word文件")
                continue
        
        try:
            wd = extract_word_text(word_path)
            word_data[proj_id] = {
                "path": str(word_path),
                "paragraph_count": wd["paragraph_count"],
                "table_count": wd["table_count"],
                "total_chars": wd["total_chars"],
                "full_text": wd["full_text"],
            }
            print(f"  ✅ {proj_id}: {wd['paragraph_count']}段, {wd['table_count']}表, {wd['total_chars']:,}字")
        except Exception as e:
            print(f"  ❌ {proj_id}: 读取失败 - {e}")
    
    # 读取JSON
    print("\n[2] 读取JSON解析结果...")
    json_data = {}
    for proj_id, base_name in PROJECT_FILE_MAP.items():
        json_path = json_dir / f"{base_name}.json"
        if not json_path.exists():
            matches = list(json_dir.glob(f"{proj_id}*.json"))
            if matches:
                json_path = matches[0]
            else:
                print(f"  ⚠️  {proj_id}: 找不到JSON文件")
                continue
        
        with open(json_path, encoding="utf-8") as f:
            blocks = json.load(f)
        
        full_text = ""
        for block in blocks:
            full_text += block.get("content", "") + "\n"
        
        json_data[proj_id] = {
            "path": str(json_path),
            "block_count": len(blocks),
            "total_chars": len(full_text),
            "full_text": full_text,
            "blocks": blocks,
        }
        print(f"  ✅ {proj_id}: {len(blocks)}块, {len(full_text):,}字")
    
    # 逐证据槽追踪
    print("\n[3] 逐题逐证据槽追踪...")
    evidence_trace = []
    question_verdicts = {}
    
    for qid, qinfo in slots_def.items():
        proj = qinfo["project"]
        task_type = qinfo["task_type"]
        
        print(f"\n  {qid} ({task_type}):")
        
        if proj not in word_data or proj not in json_data:
            print(f"    ⚠️  缺少Word或JSON数据")
            continue
        
        word_text = word_data[proj]["full_text"]
        json_text = json_data[proj]["full_text"]
        
        all_pass = True
        slot_results = []
        
        for slot in qinfo["slots"]:
            sid = slot["id"]
            fact = slot["fact"]
            keywords = slot["keywords"]
            is_internal = slot.get("internal", True)
            
            # 跳过纯外部知识槽（在第4阶段审计Web/RAG）
            if slot.get("web_only") or slot.get("rag_only"):
                slot_results.append({
                    "slot_id": sid,
                    "fact": fact,
                    "word_present": "N/A",
                    "json_present": "N/A",
                    "status": "EXTERNAL_KNOWLEDGE",
                })
                continue
            
            # 搜索Word
            word_hits = []
            for kw in keywords:
                # 标准编号用规范化搜索
                if kw.startswith("GB") or kw.startswith("HJ") or kw.startswith("DB"):
                    hits = search_standard_in_text(word_text, kw)
                    word_hits.extend(hits)
                else:
                    hits = search_in_text(word_text, [kw], context_chars=50)
                    word_hits.extend(hits)
            
            word_present = len(word_hits) > 0
            word_quote = word_hits[0]["snippet"] if word_hits else ""
            word_loc = f"char_pos:{word_hits[0]['char_pos']}" if word_hits else "NOT_FOUND"
            
            # 搜索JSON
            json_hits = []
            for kw in keywords:
                if kw.startswith("GB") or kw.startswith("HJ") or kw.startswith("DB"):
                    hits = search_standard_in_text(json_text, kw)
                    json_hits.extend(hits)
                else:
                    hits = search_in_text(json_text, [kw], context_chars=50)
                    json_hits.extend(hits)
            
            json_present = len(json_hits) > 0
            json_quote = json_hits[0]["snippet"] if json_hits else ""
            json_loc = f"char_pos:{json_hits[0]['char_pos']}" if json_hits else "NOT_FOUND"
            
            # 判断状态
            if word_present and json_present:
                status = "PASS"
            elif word_present and not json_present:
                status = "PARSE_LOSS"
                all_pass = False
            elif not word_present and json_present:
                status = "JSON_ONLY_UNEXPECTED"
            else:
                status = "SOURCE_ABSENT"
                all_pass = False
            
            slot_results.append({
                "slot_id": sid,
                "fact": fact,
                "keywords": ", ".join(keywords),
                "word_present": word_present,
                "word_quote": word_quote[:150],
                "word_location": word_loc,
                "json_present": json_present,
                "json_quote": json_quote[:150],
                "json_location": json_loc,
                "status": status,
            })
            
            if status != "PASS":
                print(f"    ⚠️  {sid} {fact}: {status}")
        
        # 计算覆盖率
        internal_slots = [s for s in slot_results if s["status"] != "EXTERNAL_KNOWLEDGE"]
        pass_count = sum(1 for s in internal_slots if s["status"] == "PASS")
        total_internal = len(internal_slots)
        
        # 题目整体状态
        if all_pass:
            verdict = "WORD_JSON_OK"
        elif pass_count > 0:
            verdict = "PARTIAL"
        else:
            verdict = "INPUT_INSUFFICIENT"
        
        question_verdicts[qid] = {
            "question_id": qid,
            "project": proj,
            "task_type": task_type,
            "total_internal_slots": total_internal,
            "pass_slots": pass_count,
            "pass_rate": f"{pass_count}/{total_internal}",
            "word_json_verdict": verdict,
            "needs_web": qinfo["needs_web"],
            "needs_rag": qinfo["needs_rag"],
        }
        
        print(f"    汇总: {pass_count}/{total_internal} 证据槽通过 → {verdict}")
        
        # 写入evidence_trace
        for sr in slot_results:
            trace_rec = {
                "question_id": qid,
                "report_id": proj,
                "evidence_slot_id": sr["slot_id"],
                "required_fact": sr["fact"],
                "keywords": sr.get("keywords", ""),
                "word_present": sr.get("word_present", "N/A"),
                "word_quote": sr.get("word_quote", ""),
                "word_location": sr.get("word_location", ""),
                "json_present": sr.get("json_present", "N/A"),
                "json_quote": sr.get("json_quote", ""),
                "json_location": sr.get("json_location", ""),
                "context_present": "PENDING",  # 阶段3补充
                "context_quote": "",
                "context_location": "",
                "status": sr["status"],
            }
            evidence_trace.append(trace_rec)
    
    # 保存evidence_trace_v2.jsonl
    trace_path = REPAIR_DIR / "evidence_trace_v2.jsonl"
    with open(trace_path, 'w', encoding='utf-8') as f:
        for rec in evidence_trace:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    print(f"\n  证据追踪已保存: {trace_path.name} ({len(evidence_trace)}条)")
    
    # 保存Word-JSON保真表
    fidelity_rows = []
    for qid, v in question_verdicts.items():
        for slot in [e for e in evidence_trace if e["question_id"] == qid]:
            fidelity_rows.append({
                "question_id": qid,
                "project": v["project"],
                "task_type": v["task_type"],
                "evidence_slot": slot["evidence_slot_id"],
                "required_fact": slot["required_fact"],
                "word_present": slot["word_present"],
                "word_location": slot["word_location"],
                "json_present": slot["json_present"],
                "json_location": slot["json_location"],
                "status": slot["status"],
            })
    
    fidelity_path = REPAIR_DIR / "02_word_json_fidelity_17.csv"
    with open(fidelity_path, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames=[
            "question_id", "project", "task_type", "evidence_slot",
            "required_fact", "word_present", "word_location",
            "json_present", "json_location", "status"
        ])
        writer.writeheader()
        writer.writerows(fidelity_rows)
    print(f"  Word-JSON保真表已保存: {fidelity_path.name} ({len(fidelity_rows)}行)")
    
    # 统计汇总
    print(f"\n[4] 17题Word-JSON保真汇总:")
    ready_count = sum(1 for v in question_verdicts.values() if v["word_json_verdict"] == "WORD_JSON_OK")
    partial_count = sum(1 for v in question_verdicts.values() if v["word_json_verdict"] == "PARTIAL")
    insufficient_count = sum(1 for v in question_verdicts.values() if v["word_json_verdict"] == "INPUT_INSUFFICIENT")
    
    print(f"  全部通过(WORD_JSON_OK): {ready_count}/17")
    print(f"  部分通过(PARTIAL): {partial_count}/17")
    print(f"  输入不足(INPUT_INSUFFICIENT): {insufficient_count}/17")
    
    for qid, v in sorted(question_verdicts.items()):
        status_icon = "✅" if v["word_json_verdict"] == "WORD_JSON_OK" else "⚠️"
        print(f"    {status_icon} {qid}: {v['pass_rate']} → {v['word_json_verdict']}")
    
    # 保存逐题裁决
    verdict_path = REPAIR_DIR / "08_question_verdicts_summary_v2.csv"
    with open(verdict_path, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames=[
            "question_id", "project", "task_type", "total_internal_slots",
            "pass_slots", "pass_rate", "word_json_verdict",
            "needs_web", "needs_rag", "overall_status"
        ])
        writer.writeheader()
        for qid, v in sorted(question_verdicts.items()):
            v["overall_status"] = "PENDING_CONTEXT_AUDIT"  # 待阶段3补充
            writer.writerow(v)
    print(f"  逐题裁决表已保存: {verdict_path.name}")
    
    print(f"\n阶段2完成 (Word→JSON部分)")
    print(f"  注意：上下文证据槽审计将在阶段3补充")
    return evidence_trace, question_verdicts

if __name__ == "__main__":
    main()
